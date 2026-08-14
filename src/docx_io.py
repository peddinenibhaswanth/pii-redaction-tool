"""
docx_io.py
----------
DOCX reading/writing utilities for the PII redaction pipeline.

Handles:
- body paragraphs
- tables and nested tables
- headers and footers
- text boxes / drawing shapes in the main document
- text boxes / drawing shapes in headers and footers
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def _iter_paragraphs_in_table(table):
    """Yield paragraphs in a table, including nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph

            for nested in cell.tables:
                yield from _iter_paragraphs_in_table(nested)


def _get_xml_element(parent):
    """
    Get the underlying XML element from a python-docx object.

    Document exposes .element, while Header/Footer objects expose
    ._element.
    """

    if hasattr(parent, "element"):
        return parent.element

    if hasattr(parent, "_element"):
        return parent._element

    raise TypeError(
        f"Unsupported DOCX parent object: {type(parent).__name__}"
    )


def _iter_textbox_paragraphs(parent):
    """
    Yield Paragraph objects inside Word text boxes / drawing shapes.

    python-docx does not expose these paragraphs through the normal
    .paragraphs property, so we inspect the underlying XML.
    """

    root = _get_xml_element(parent)

    for textbox in root.iter(qn("w:txbxContent")):

        for paragraph_element in textbox.iter(qn("w:p")):

            yield Paragraph(
                paragraph_element,
                parent
            )


def _iter_section_part_paragraphs(part):
    """Yield normal and table paragraphs from a header/footer part."""

    for paragraph in part.paragraphs:
        yield paragraph

    for table in part.tables:
        yield from _iter_paragraphs_in_table(table)


def iter_all_paragraphs(doc: Document):
    """
    Yield every visible paragraph that we can process.

    Covered structures:
      - body paragraphs
      - body tables
      - nested tables
      - headers
      - footers
      - first-page headers/footers
      - even-page headers/footers
      - text boxes in the main document
      - text boxes in headers/footers
    """

    # ---------------------------------------------------------
    # Main document body
    # ---------------------------------------------------------

    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        yield from _iter_paragraphs_in_table(table)

    # Text boxes / shapes in the main document.
    yield from _iter_textbox_paragraphs(doc)

    # ---------------------------------------------------------
    # Headers and footers
    # ---------------------------------------------------------

    for section in doc.sections:

        parts = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )

        for part in parts:

            if part is None:
                continue

            yield from _iter_section_part_paragraphs(part)

            # Text boxes / shapes inside this header/footer.
            yield from _iter_textbox_paragraphs(part)


def get_paragraph_text(paragraph):
    """Return the visible text represented by a paragraph's runs."""

    return "".join(
        run.text
        for run in paragraph.runs
    )


def set_paragraph_text(paragraph, new_text):
    """
    Replace a paragraph's visible text with new_text.

    The first run receives the complete replacement text and subsequent
    runs are emptied.

    This preserves paragraph-level formatting and structure, but does
    not perfectly preserve formatting that changes within a paragraph.
    """

    if not paragraph.runs:
        return

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


def load_document(path):
    """Load a DOCX document."""

    return Document(path)


def save_document(doc, path):
    """Save a DOCX document."""

    doc.save(path)